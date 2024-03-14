""" from docx import Document

import re



def docx_replace_regex(doc_obj, regex , replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

def getText(doc_obj):
    fullText = []
    for para in doc_obj.paragraphs:
        fullText.append(para.text)
    return '\\n'.join(fullText)

doc = Document('C:\\Users\\umesh.bihani\\Documents\\Ontab\\docs\\Pyaout_spuer\\ADHOC MF BURN Delinquent email_v4.0')
print(getText(doc)) """
import aspose.words as aw


doc = aw.Document('C:\\Users\\umesh.bihani\\Documents\\Ontab\\docs\\Pyaout_spuer\\ADHOC MF-Payout Delinquent email_v4.0.docx')

doc.range.replace("#PAYOUT_AMOUNT", "Hello", 
    aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))

doc.
